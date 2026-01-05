"""Service for generating COA documents."""

import os
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any, List
from io import BytesIO
import zipfile

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, 
    Image, PageBreak, KeepTogether
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from sqlalchemy.orm import Session
from loguru import logger

from ..models import Lot, LotStatus
from ..models.coa import COAHistory
from ..services.base import BaseService
from ..config import settings


class COAGeneratorService:
    """Service for generating COA documents from approved lots."""

    def __init__(self):
        self.template_dir = Path("templates")
        self.output_dir = Path(settings.COA_OUTPUT_FOLDER)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_coa(
        self,
        db: Session,
        lot_id: int,
        template: str = "standard",
        output_format: str = "pdf",
        user_id: Optional[int] = None,
    ) -> Dict[str, Any]:
        """Generate COA for a lot."""
        lot = db.query(Lot).filter(Lot.id == lot_id).first()
        if not lot:
            raise ValueError(f"Lot {lot_id} not found")

        if lot.status not in [LotStatus.APPROVED, LotStatus.RELEASED]:
            raise ValueError(f"Lot {lot.lot_number} is not approved for COA generation (status: {lot.status.value})")

        if not lot.generate_coa:
            raise ValueError(f"Lot {lot.lot_number} is not marked for COA generation")

        # Check if all test results are approved
        from ..models.enums import TestResultStatus
        unapproved = [r for r in lot.test_results if r.status != TestResultStatus.APPROVED]
        if unapproved:
            raise ValueError(f"Lot has {len(unapproved)} unapproved test results")

        try:
            # Generate filename
            filename_base = self._generate_filename(lot)

            # Generate documents
            generated_files = []

            if output_format in ["docx", "both"]:
                docx_path = self._generate_docx(lot, template, filename_base)
                generated_files.append(docx_path)

            if output_format in ["pdf", "both"]:
                pdf_path = self._generate_pdf(lot, template, filename_base)
                generated_files.append(pdf_path)

            # Update lot status (only if not already released)
            if lot.status != LotStatus.RELEASED:
                lot.status = LotStatus.RELEASED

            # Create COA history entry
            for file_path in generated_files:
                coa_history = COAHistory(
                    lot_id=lot_id,
                    filename=file_path.name,
                    generated_by=str(user_id) if user_id else "system",
                )
                db.add(coa_history)

            db.commit()

            logger.info(f"Generated COA for lot {lot.lot_number}")

            return {
                "status": "success",
                "files": generated_files,
                "lot_number": lot.lot_number,
            }

        except Exception as e:
            logger.error(f"Failed to generate COA: {e}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            db.rollback()
            raise

    def _generate_filename(self, lot: Lot) -> str:
        """Generate standardized filename for COA."""
        date_str = datetime.now().strftime("%Y%m%d")

        # Get primary product for naming
        primary_product = lot.lot_products[0].product if lot.lot_products else None
        if primary_product:
            brand = primary_product.brand.replace(" ", "")
            product = primary_product.product_name.replace(" ", "")
        else:
            brand = "Unknown"
            product = "Product"

        lot_number = lot.lot_number.replace(" ", "")

        return f"{date_str}-{brand}-{product}-{lot_number}"

    def _generate_docx(self, lot: Lot, template: str, filename_base: str) -> Path:
        """Generate Word document COA."""
        # Create new document
        doc = Document()

        # Add header
        self._add_docx_header(doc, lot)

        # Add product information
        self._add_product_info_docx(doc, lot)

        # Add test results
        self._add_test_results_docx(doc, lot)

        # Add footer
        self._add_docx_footer(doc, lot)

        # Save document
        output_path = self.output_dir / f"{filename_base}.docx"
        doc.save(output_path)

        return output_path

    def _add_docx_header(self, doc: Document, lot: Lot):
        """Add header to DOCX document."""
        # Title
        title = doc.add_heading("Certificate of Analysis", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Company info
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Your Company Name\n").bold = True
        p.add_run("123 Quality Street, Lab City, LC 12345\n")
        p.add_run("Phone: (555) 123-4567 | Email: lab@company.com")

        doc.add_paragraph()  # Spacing

    def _add_product_info_docx(self, doc: Document, lot: Lot):
        """Add product information section to DOCX."""
        doc.add_heading("Product Information", level=1)

        # Create info table
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light List"

        # Add rows
        info_items = [
            ("Reference Number:", lot.reference_number),
            ("Lot Number:", lot.lot_number),
            (
                "Manufacturing Date:",
                lot.mfg_date.strftime("%B %d, %Y") if lot.mfg_date else "N/A",
            ),
            (
                "Expiration Date:",
                lot.exp_date.strftime("%B %d, %Y") if lot.exp_date else "N/A",
            ),
            ("Release Date:", datetime.now().strftime("%B %d, %Y")),
        ]

        # Add product(s)
        if lot.lot_type.value == "multi_sku_composite":
            products_str = "\n".join(
                [
                    f"• {lp.product.display_name} ({lp.percentage}%)"
                    for lp in lot.lot_products
                ]
            )
            info_items.append(("Products:", products_str))
        else:
            product = lot.lot_products[0].product if lot.lot_products else None
            if product:
                info_items.extend(
                    [
                        ("Brand:", product.brand),
                        ("Product:", product.product_name),
                        ("Flavor:", product.flavor or "N/A"),
                        ("Size:", product.size or "N/A"),
                    ]
                )

        for label, value in info_items:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = str(value)

        doc.add_paragraph()  # Spacing

    def _add_test_results_docx(self, doc: Document, lot: Lot):
        """Add test results section to DOCX."""
        doc.add_heading("Test Results", level=1)

        # Group test results by category
        microbiological = []
        heavy_metals = []
        other = []

        for result in lot.test_results:
            if result.test_type in [
                "Total Plate Count",
                "Yeast/Mold",
                "E. Coli",
                "Salmonella",
                "Staphylococcus aureus",
                "Total Coliform Count",
            ]:
                microbiological.append(result)
            elif result.test_type in ["Lead", "Mercury", "Cadmium", "Arsenic"]:
                heavy_metals.append(result)
            else:
                other.append(result)

        # Microbiological results
        if microbiological:
            doc.add_heading("Microbiological Analysis", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid"

            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = "Test Parameter"
            header_cells[1].text = "Result"
            header_cells[2].text = "Specification"

            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True

            # Data rows
            for result in microbiological:
                row = table.add_row()
                row.cells[0].text = result.test_type
                row.cells[1].text = f"{result.result_value} {result.unit}".strip()
                row.cells[2].text = self._get_specification(result.test_type)

        # Heavy metals
        if heavy_metals:
            doc.add_paragraph()  # Spacing
            doc.add_heading("Heavy Metals Analysis", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid"

            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = "Test Parameter"
            header_cells[1].text = "Result"
            header_cells[2].text = "Specification"

            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True

            # Data rows
            for result in heavy_metals:
                row = table.add_row()
                row.cells[0].text = result.test_type
                row.cells[1].text = f"{result.result_value} {result.unit}".strip()
                row.cells[2].text = self._get_specification(result.test_type)

        # Other tests
        if other:
            doc.add_paragraph()  # Spacing
            doc.add_heading("Additional Tests", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid"

            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = "Test Parameter"
            header_cells[1].text = "Result"
            header_cells[2].text = "Specification"

            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True

            # Data rows
            for result in other:
                row = table.add_row()
                row.cells[0].text = result.test_type
                row.cells[1].text = f"{result.result_value} {result.unit}".strip()
                row.cells[2].text = self._get_specification(result.test_type)

        doc.add_paragraph()  # Spacing

    def _add_docx_footer(self, doc: Document, lot: Lot):
        """Add footer to DOCX document."""
        # Approval section
        doc.add_heading("Quality Assurance", level=1)

        p = doc.add_paragraph()
        p.add_run(
            "This product has been tested and released in accordance with established specifications.\n\n"
        )

        # Signature line
        table = doc.add_table(rows=1, cols=2)

        # QC Approval
        cell1 = table.rows[0].cells[0]
        cell1.text = "QC Approved By: _________________\n\n"
        approver_name = "QC Manager"
        if lot.test_results and lot.test_results[0].approved_by_user:
            approver_name = lot.test_results[0].approved_by_user.username
        cell1.text += f"Name: {approver_name}\n"
        cell1.text += f"Date: {datetime.now().strftime('%B %d, %Y')}"

        # Authorized Signature
        cell2 = table.rows[0].cells[1]
        cell2.text = "Authorized Signature: _________________\n\n"
        cell2.text += "Name: _________________\n"
        cell2.text += f"Date: {datetime.now().strftime('%B %d, %Y')}"

        # Disclaimer
        doc.add_paragraph()
        disclaimer = doc.add_paragraph()
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        disclaimer.add_run(
            "This Certificate of Analysis applies only to the specific lot referenced above. "
            "Results relate only to the items tested."
        ).font.size = Pt(8)

    def _generate_pdf(self, lot: Lot, template: str, filename_base: str) -> Path:
        """Generate PDF COA using ReportLab."""
        output_path = self.output_dir / f"{filename_base}.pdf"
        
        try:
            # Create PDF directly to file, not buffer
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=letter,
                rightMargin=0.4*inch,
                leftMargin=0.4*inch,
                topMargin=0.5*inch,
                bottomMargin=0.5*inch
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            self._setup_custom_styles(styles)
            
            # Build story
            story = []
            
            # Add header
            story.extend(self._create_pdf_header(styles))
            
            # Add title
            story.append(Paragraph("CERTIFICATE OF ANALYSIS", styles['COATitle']))
            story.append(Spacer(1, 0.15*inch))
            
            # Add lot information
            story.extend(self._create_pdf_lot_info(lot, styles))
            story.append(Spacer(1, 0.15*inch))
            
            # Add test results
            story.extend(self._create_pdf_test_results(lot, styles))
            story.append(Spacer(1, 0.15*inch))
            
            # Add certification
            story.extend(self._create_pdf_certification(styles))
            story.append(Spacer(1, 0.15*inch))
            
            # Add signatures
            story.extend(self._create_pdf_signatures(lot, styles))
            
            # Build PDF
            doc.build(story)
            
            logger.info(f"PDF generated successfully at {output_path}")
            logger.info(f"PDF file size: {output_path.stat().st_size} bytes")
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating PDF: {e}")
            import traceback
            logger.error(f"PDF generation traceback: {traceback.format_exc()}")
            raise
    
    def _setup_custom_styles(self, styles):
        """Setup custom paragraph styles for PDF."""
        # Title style
        styles.add(ParagraphStyle(
            name='COATitle',
            parent=styles['Title'],
            fontSize=18,
            textColor=colors.HexColor('#1f4788'),
            alignment=TA_CENTER,
            spaceAfter=10
        ))
        
        # Header style
        styles.add(ParagraphStyle(
            name='COAHeader',
            parent=styles['Heading1'],
            fontSize=12,
            textColor=colors.HexColor('#1f4788'),
            alignment=TA_LEFT,
            spaceAfter=6
        ))
        
        # Normal text
        styles.add(ParagraphStyle(
            name='COANormal',
            parent=styles['Normal'],
            fontSize=9,
            alignment=TA_LEFT,
            leading=10
        ))
        
        # Footer style
        styles.add(ParagraphStyle(
            name='COAFooter',
            parent=styles['Normal'],
            fontSize=8,
            alignment=TA_CENTER,
            textColor=colors.grey
        ))
    
    def _create_pdf_header(self, styles) -> List:
        """Create PDF header."""
        header_data = [
            ['Body Nutrition', '', 'Tel: (727) 555-0123'],
            ['2950 47th Ave North', '', 'Email: quality@bodynutrition.com'],
            ['St Petersburg, FL 33714', '', 'www.bodynutrition.com']
        ]
        
        header_table = Table(header_data, colWidths=[4*inch, 2*inch, 2*inch])
        header_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),
            ('ALIGN', (2, 0), (2, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (0, 0), 14),
            ('TEXTCOLOR', (0, 0), (0, 0), colors.HexColor('#1f4788')),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.grey),
        ]))
        
        return [header_table, Spacer(1, 0.2*inch)]
    
    def _create_pdf_lot_info(self, lot: Lot, styles) -> List:
        """Create lot information section for PDF."""
        elements = []
        
        # Get product names
        products = ", ".join([lp.product.display_name for lp in lot.lot_products])
        
        # Create lot info table
        lot_data = [
            ['Product:', products],
            ['Lot Number:', lot.lot_number],
            ['Reference Number:', lot.reference_number],
            ['Manufacture Date:', lot.mfg_date.strftime('%B %d, %Y') if lot.mfg_date else 'N/A'],
            ['Expiration Date:', lot.exp_date.strftime('%B %d, %Y') if lot.exp_date else 'N/A'],
        ]
        
        lot_table = Table(lot_data, colWidths=[2*inch, 5*inch])
        lot_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 2),
        ]))
        
        elements.append(Paragraph("PRODUCT INFORMATION", styles['COAHeader']))
        elements.append(lot_table)
        
        return elements
    
    def _create_pdf_test_results(self, lot: Lot, styles) -> List:
        """Create test results section for PDF."""
        elements = []
        
        elements.append(Paragraph("TEST RESULTS", styles['COAHeader']))
        
        # Group test results by category
        microbiological = []
        heavy_metals = []
        other = []
        
        for result in lot.test_results:
            if "plate count" in result.test_type.lower() or \
               "yeast" in result.test_type.lower() or \
               "mold" in result.test_type.lower() or \
               "coli" in result.test_type.lower() or \
               "salmonella" in result.test_type.lower():
                microbiological.append(result)
            elif any(metal in result.test_type.lower() for metal in ["lead", "mercury", "cadmium", "arsenic"]):
                heavy_metals.append(result)
            else:
                other.append(result)
        
        # Create tables for each category
        if microbiological:
            elements.append(Spacer(1, 0.05*inch))
            elements.append(Paragraph("<b>Microbiological Analysis</b>", styles['COANormal']))
            elements.append(Spacer(1, 0.1*inch))
            
            table_data = [['Test Parameter', 'Result', 'Unit', 'Specification']]
            for result in microbiological:
                table_data.append([
                    result.test_type,
                    result.result_value or 'ND',
                    result.unit or '',
                    result.specification or self._get_specification(result.test_type)
                ])
            
            table = Table(table_data, colWidths=[2.5*inch, 1.5*inch, 1*inch, 2*inch])
            table.setStyle(TableStyle([
                # Header row
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e8f0fe')),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                
                # Data rows
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('ALIGN', (1, 1), (2, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                
                # Grid
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('LINEBELOW', (0, 0), (-1, 0), 1, colors.black),
                
                # Padding
                ('TOPPADDING', (0, 0), (-1, -1), 3),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ]))
            
            elements.append(table)
        
        if heavy_metals:
            elements.append(Spacer(1, 0.05*inch))
            elements.append(Paragraph("<b>Heavy Metals Analysis</b>", styles['COANormal']))
            elements.append(Spacer(1, 0.1*inch))
            
            table_data = [['Test Parameter', 'Result', 'Unit', 'Specification']]
            for result in heavy_metals:
                table_data.append([
                    result.test_type,
                    result.result_value or 'ND',
                    result.unit or '',
                    result.specification or self._get_specification(result.test_type)
                ])
            
            table = Table(table_data, colWidths=[2.5*inch, 1.5*inch, 1*inch, 2*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e8f0fe')),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('ALIGN', (1, 1), (2, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('LINEBELOW', (0, 0), (-1, 0), 1, colors.black),
                ('TOPPADDING', (0, 0), (-1, -1), 3),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ]))
            
            elements.append(table)
        
        if other:
            elements.append(Spacer(1, 0.05*inch))
            elements.append(Paragraph("<b>Additional Tests</b>", styles['COANormal']))
            elements.append(Spacer(1, 0.1*inch))
            
            table_data = [['Test Parameter', 'Result', 'Unit', 'Specification']]
            for result in other:
                table_data.append([
                    result.test_type,
                    result.result_value or 'ND',
                    result.unit or '',
                    result.specification or self._get_specification(result.test_type)
                ])
            
            table = Table(table_data, colWidths=[2.5*inch, 1.5*inch, 1*inch, 2*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e8f0fe')),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('ALIGN', (1, 1), (2, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('LINEBELOW', (0, 0), (-1, 0), 1, colors.black),
                ('TOPPADDING', (0, 0), (-1, -1), 3),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ]))
            
            elements.append(table)
        
        return elements
    
    def _create_pdf_certification(self, styles) -> List:
        """Create certification section for PDF."""
        elements = []
        
        cert_text = """The above results apply to the sample as received. This Certificate of Analysis is not to be reproduced except in full, without written approval from the laboratory. This product has been tested and meets all specifications for release."""
        
        elements.append(Paragraph("CERTIFICATION", styles['COAHeader']))
        elements.append(Paragraph(cert_text, styles['COANormal']))
        
        return elements
    
    def _create_pdf_signatures(self, lot: Lot, styles) -> List:
        """Create signature section for PDF."""
        elements = []
        
        # Get approver info
        approver_name = "Quality Control Manager"
        if lot.test_results and lot.test_results[0].approved_by_user:
            approver_name = lot.test_results[0].approved_by_user.username
        
        # Create signature table with more space
        sig_data = [
            ['', ''],  # Empty row for signatures
            ['_' * 40, '_' * 40],
            [approver_name, 'Quality Assurance Director'],
            ['Date: ' + datetime.now().strftime('%m/%d/%Y'), 'Date: _____________'],
        ]
        
        sig_table = Table(sig_data, colWidths=[3.5*inch, 3.5*inch])
        sig_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTSIZE', (0, 2), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (0, 0), 15),  # More space above signature line
            ('BOTTOMPADDING', (0, 1), (0, 1), 5),  # Space below signature line
            ('TOPPADDING', (0, 2), (-1, -1), 2),
            ('BOTTOMPADDING', (0, 2), (-1, -1), 2),
        ]))
        
        elements.append(sig_table)
        
        # Footer
        elements.append(Spacer(1, 0.1*inch))
        footer_text = f"Generated on {datetime.now().strftime('%m/%d/%Y at %I:%M %p')}"
        elements.append(Paragraph(footer_text, styles['COAFooter']))
        
        return elements

    def _get_specification(self, test_type: str) -> str:
        """Get specification limits for test types."""
        specs = {
            "Total Plate Count": "< 10,000 CFU/g",
            "Yeast/Mold": "< 1,000 CFU/g",
            "E. Coli": "Negative",
            "Salmonella": "Negative",
            "Staphylococcus aureus": "Negative",
            "Total Coliform Count": "< 10 CFU/g",
            "Lead": "< 0.5 ppm",
            "Mercury": "< 0.1 ppm",
            "Cadmium": "< 0.3 ppm",
            "Arsenic": "< 1.0 ppm",
            "Gluten": "< 20 ppm",
        }
        return specs.get(test_type, "Within limits")

    def generate_batch_coas(
        self,
        db: Session,
        lot_ids: List[int],
        template: str = "standard",
        output_format: str = "pdf",
        user_id: Optional[int] = None,
    ) -> Dict[str, Any]:
        """Generate COAs for multiple lots."""
        results = {"success": [], "failed": [], "files": []}

        for lot_id in lot_ids:
            try:
                result = self.generate_coa(db, lot_id, template, output_format, user_id)
                results["success"].append(lot_id)
                results["files"].extend(result["files"])
            except Exception as e:
                logger.error(f"Failed to generate COA for lot {lot_id}: {e}")
                results["failed"].append({"lot_id": lot_id, "error": str(e)})

        # Create ZIP file if multiple files
        if len(results["files"]) > 1:
            zip_path = self._create_zip(results["files"])
            results["zip_file"] = zip_path

        return results

    def _create_zip(self, files: List[Path]) -> Path:
        """Create ZIP file containing multiple COAs."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        zip_path = self.output_dir / f"COAs_batch_{timestamp}.zip"

        with zipfile.ZipFile(zip_path, "w") as zf:
            for file_path in files:
                if file_path.exists():
                    zf.write(file_path, file_path.name)

        return zip_path

    def get_coa_history(self, lot_id: Optional[int] = None) -> List[COAHistory]:
        """Get COA generation history."""
        query = self.session.query(COAHistory)
        if lot_id:
            query = query.filter_by(lot_id=lot_id)
        return query.order_by(COAHistory.generated_at.desc()).all()
