"""Service for generating COA documents."""

import os
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any, List
from io import BytesIO
import zipfile

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from sqlalchemy.orm import Session
from loguru import logger

from ..models import Lot, COAHistory, LotStatus
from ..services.base import BaseService
from ..config import settings


class COAGeneratorService:
    """Service for generating COA documents from approved lots."""

    def __init__(self):
        self.template_dir = Path("templates")
        self.output_dir = Path(settings.COA_OUTPUT_FOLDER)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_coa(
        self,
        db: Session,
        lot_id: int,
        template: str = "standard",
        output_format: str = "pdf",
        user_id: Optional[int] = None,
    ) -> Dict[str, Any]:
        """Generate COA for a lot."""
        lot = db.query(Lot).filter(Lot.id == lot_id).first()
        if not lot:
            raise ValueError(f"Lot {lot_id} not found")

        if lot.status != LotStatus.APPROVED:
            raise ValueError(f"Lot {lot.lot_number} is not approved for COA generation")

        if not lot.generate_coa:
            raise ValueError(f"Lot {lot.lot_number} is not marked for COA generation")

        # Check if all test results are approved
        unapproved = [r for r in lot.test_results if r.status != "approved"]
        if unapproved:
            raise ValueError(f"Lot has {len(unapproved)} unapproved test results")

        try:
            # Generate filename
            filename_base = self._generate_filename(lot)

            # Generate documents
            generated_files = []

            if output_format in ["docx", "both"]:
                docx_path = self._generate_docx(lot, template, filename_base)
                generated_files.append(docx_path)

            if output_format in ["pdf", "both"]:
                pdf_path = self._generate_pdf(lot, template, filename_base)
                generated_files.append(pdf_path)

            # Update lot status
            lot.status = LotStatus.RELEASED

            # Create COA history entry
            for file_path in generated_files:
                coa_history = COAHistory(
                    lot_id=lot_id,
                    filename=file_path.name,
                    generated_by=str(user_id) if user_id else "system",
                )
                self.session.add(coa_history)

            self.session.commit()

            logger.info(f"Generated COA for lot {lot.lot_number}")

            return {
                "status": "success",
                "files": generated_files,
                "lot_number": lot.lot_number,
            }

        except Exception as e:
            logger.error(f"Failed to generate COA: {e}")
            self.session.rollback()
            raise

    def _generate_filename(self, lot: Lot) -> str:
        """Generate standardized filename for COA."""
        date_str = datetime.now().strftime("%Y%m%d")

        # Get primary product for naming
        primary_product = lot.lot_products[0].product if lot.lot_products else None
        if primary_product:
            brand = primary_product.brand.replace(" ", "")
            product = primary_product.product_name.replace(" ", "")
        else:
            brand = "Unknown"
            product = "Product"

        lot_number = lot.lot_number.replace(" ", "")

        return f"{date_str}-{brand}-{product}-{lot_number}"

    def _generate_docx(self, lot: Lot, template: str, filename_base: str) -> Path:
        """Generate Word document COA."""
        # Create new document
        doc = Document()

        # Add header
        self._add_docx_header(doc, lot)

        # Add product information
        self._add_product_info_docx(doc, lot)

        # Add test results
        self._add_test_results_docx(doc, lot)

        # Add footer
        self._add_docx_footer(doc, lot)

        # Save document
        output_path = self.output_dir / f"{filename_base}.docx"
        doc.save(output_path)

        return output_path

    def _add_docx_header(self, doc: Document, lot: Lot):
        """Add header to DOCX document."""
        # Title
        title = doc.add_heading("Certificate of Analysis", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Company info
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Your Company Name\n").bold = True
        p.add_run("123 Quality Street, Lab City, LC 12345\n")
        p.add_run("Phone: (555) 123-4567 | Email: lab@company.com")

        doc.add_paragraph()  # Spacing

    def _add_product_info_docx(self, doc: Document, lot: Lot):
        """Add product information section to DOCX."""
        doc.add_heading("Product Information", level=1)

        # Create info table
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light List"

        # Add rows
        info_items = [
            ("Reference Number:", lot.reference_number),
            ("Lot Number:", lot.lot_number),
            (
                "Manufacturing Date:",
                lot.mfg_date.strftime("%B %d, %Y") if lot.mfg_date else "N/A",
            ),
            (
                "Expiration Date:",
                lot.exp_date.strftime("%B %d, %Y") if lot.exp_date else "N/A",
            ),
            ("Release Date:", datetime.now().strftime("%B %d, %Y")),
        ]

        # Add product(s)
        if lot.lot_type.value == "multi_sku_composite":
            products_str = "\n".join(
                [
                    f"• {lp.product.display_name} ({lp.percentage}%)"
                    for lp in lot.lot_products
                ]
            )
            info_items.append(("Products:", products_str))
        else:
            product = lot.lot_products[0].product if lot.lot_products else None
            if product:
                info_items.extend(
                    [
                        ("Brand:", product.brand),
                        ("Product:", product.product_name),
                        ("Flavor:", product.flavor or "N/A"),
                        ("Size:", product.size or "N/A"),
                    ]
                )

        for label, value in info_items:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = str(value)

        doc.add_paragraph()  # Spacing

    def _add_test_results_docx(self, doc: Document, lot: Lot):
        """Add test results section to DOCX."""
        doc.add_heading("Test Results", level=1)

        # Group test results by category
        microbiological = []
        heavy_metals = []
        other = []

        for result in lot.test_results:
            if result.test_type in [
                "Total Plate Count",
                "Yeast/Mold",
                "E. Coli",
                "Salmonella",
                "Staphylococcus aureus",
                "Total Coliform Count",
            ]:
                microbiological.append(result)
            elif result.test_type in ["Lead", "Mercury", "Cadmium", "Arsenic"]:
                heavy_metals.append(result)
            else:
                other.append(result)

        # Microbiological results
        if microbiological:
            doc.add_heading("Microbiological Analysis", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid"

            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = "Test Parameter"
            header_cells[1].text = "Result"
            header_cells[2].text = "Specification"

            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True

            # Data rows
            for result in microbiological:
                row = table.add_row()
                row.cells[0].text = result.test_type
                row.cells[1].text = f"{result.result_value} {result.unit}".strip()
                row.cells[2].text = self._get_specification(result.test_type)

        # Heavy metals
        if heavy_metals:
            doc.add_paragraph()  # Spacing
            doc.add_heading("Heavy Metals Analysis", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid"

            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = "Test Parameter"
            header_cells[1].text = "Result"
            header_cells[2].text = "Specification"

            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True

            # Data rows
            for result in heavy_metals:
                row = table.add_row()
                row.cells[0].text = result.test_type
                row.cells[1].text = f"{result.result_value} {result.unit}".strip()
                row.cells[2].text = self._get_specification(result.test_type)

        # Other tests
        if other:
            doc.add_paragraph()  # Spacing
            doc.add_heading("Additional Tests", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid"

            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = "Test Parameter"
            header_cells[1].text = "Result"
            header_cells[2].text = "Specification"

            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True

            # Data rows
            for result in other:
                row = table.add_row()
                row.cells[0].text = result.test_type
                row.cells[1].text = f"{result.result_value} {result.unit}".strip()
                row.cells[2].text = self._get_specification(result.test_type)

        doc.add_paragraph()  # Spacing

    def _add_docx_footer(self, doc: Document, lot: Lot):
        """Add footer to DOCX document."""
        # Approval section
        doc.add_heading("Quality Assurance", level=1)

        p = doc.add_paragraph()
        p.add_run(
            "This product has been tested and released in accordance with established specifications.\n\n"
        )

        # Signature line
        table = doc.add_table(rows=1, cols=2)

        # QC Approval
        cell1 = table.rows[0].cells[0]
        cell1.text = "QC Approved By: _________________\n\n"
        cell1.text += f"Name: {lot.test_results[0].approved_by if lot.test_results else 'QC Manager'}\n"
        cell1.text += f"Date: {datetime.now().strftime('%B %d, %Y')}"

        # Authorized Signature
        cell2 = table.rows[0].cells[1]
        cell2.text = "Authorized Signature: _________________\n\n"
        cell2.text += "Name: _________________\n"
        cell2.text += f"Date: {datetime.now().strftime('%B %d, %Y')}"

        # Disclaimer
        doc.add_paragraph()
        disclaimer = doc.add_paragraph()
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        disclaimer.add_run(
            "This Certificate of Analysis applies only to the specific lot referenced above. "
            "Results relate only to the items tested."
        ).font.size = Pt(8)

    def _generate_pdf(self, lot: Lot, template: str, filename_base: str) -> Path:
        """Generate PDF COA (simplified version)."""
        output_path = self.output_dir / f"{filename_base}.pdf"

        # For now, generate DOCX first then convert
        # In production, you would use ReportLab or convert DOCX to PDF
        docx_path = self.output_dir / f"{filename_base}.docx"
        if not docx_path.exists():
            self._generate_docx(lot, template, filename_base)

        # TODO: Implement proper PDF generation or conversion
        # For now, just copy the path
        return docx_path.with_suffix(".pdf")

    def _get_specification(self, test_type: str) -> str:
        """Get specification limits for test types."""
        specs = {
            "Total Plate Count": "< 10,000 CFU/g",
            "Yeast/Mold": "< 1,000 CFU/g",
            "E. Coli": "Negative",
            "Salmonella": "Negative",
            "Staphylococcus aureus": "Negative",
            "Total Coliform Count": "< 10 CFU/g",
            "Lead": "< 0.5 ppm",
            "Mercury": "< 0.1 ppm",
            "Cadmium": "< 0.3 ppm",
            "Arsenic": "< 1.0 ppm",
            "Gluten": "< 20 ppm",
        }
        return specs.get(test_type, "Within limits")

    def generate_batch_coas(
        self,
        lot_ids: List[int],
        template: str = "standard",
        output_format: str = "pdf",
        user_id: Optional[int] = None,
    ) -> Dict[str, Any]:
        """Generate COAs for multiple lots."""
        results = {"success": [], "failed": [], "files": []}

        for lot_id in lot_ids:
            try:
                result = self.generate_coa(lot_id, template, output_format, user_id)
                results["success"].append(lot_id)
                results["files"].extend(result["files"])
            except Exception as e:
                logger.error(f"Failed to generate COA for lot {lot_id}: {e}")
                results["failed"].append({"lot_id": lot_id, "error": str(e)})

        # Create ZIP file if multiple files
        if len(results["files"]) > 1:
            zip_path = self._create_zip(results["files"])
            results["zip_file"] = zip_path

        return results

    def _create_zip(self, files: List[Path]) -> Path:
        """Create ZIP file containing multiple COAs."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        zip_path = self.output_dir / f"COAs_batch_{timestamp}.zip"

        with zipfile.ZipFile(zip_path, "w") as zf:
            for file_path in files:
                if file_path.exists():
                    zf.write(file_path, file_path.name)

        return zip_path

    def get_coa_history(self, lot_id: Optional[int] = None) -> List[COAHistory]:
        """Get COA generation history."""
        query = self.session.query(COAHistory)
        if lot_id:
            query = query.filter_by(lot_id=lot_id)
        return query.order_by(COAHistory.generated_at.desc()).all()
