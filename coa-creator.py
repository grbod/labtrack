import pandas as pd
from docx import Document
from datetime import datetime
import os
# import win32com.client  # Not available on macOS
# import pythoncom  # Not available on macOS
import sys


def format_date(date_value):
    if isinstance(date_value, datetime):
        return date_value.strftime('%Y-%m-%d')
    return str(date_value)

def fill_template(row, template_path):
    doc = Document(template_path)
    
    placeholders = {
        '<Brand>': row['Brand'],
        '<Product>': row['Product'],
        '<Flavor>': row['Flavor'],
        '<Size>': row['Size'],
        '<Lot>': row['Lot'],
        '<MfgDate>': format_date(row['Mfg Date']),
        '<ExpDate>': format_date(row['Exp Date']),
        '<refID>': row['RefID'],
        '<PlateCount>': row['Total Plate Count'],
        '<YeastMold>': row['Yeast/Mold'],
        '<EColi>': row['E. Coli'],
        '<Salmonella>': row['Salmonella'],
        '<ReleaseDate>': format_date(row['Release Date']),
        '<QCname>': row['QC Approval'],
        '<Staphylococcus>': row['Staphylococcus aureus'],
        '<Coliform>': row['Total Coliform Count'],
        '<Gluten>': row['Gluten'],
        '<Lead>': row['Lead'],
        '<Mercury>': row['Mercury'],
        '<Cadmium>': row['Cadmium'],
        '<Arsenic>': row['Arsenic']
    }

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, '' if pd.isna(value) else str(value))

    # Handle tables
    for table in doc.tables:
        rows_to_delete = []
        for i, table_row in enumerate(table.rows):
            row_has_nan = False
            for cell in table_row.cells:
                for placeholder, value in placeholders.items():
                    if placeholder in cell.text:
                        if pd.isna(value):
                            row_has_nan = True
                        else:
                            cell.text = cell.text.replace(placeholder, str(value))
            if row_has_nan:
                rows_to_delete.append(i)
        
        # Delete rows with NaN values (in reverse order to avoid index issues)
        for row_index in sorted(rows_to_delete, reverse=True):
            table._element.remove(table.rows[row_index]._element)

    return doc

def save_as_docx_and_pdf(doc, filename_base, output_directory):
    # Save the Word document
    docx_path = os.path.join(output_directory, f"{filename_base}.docx")
    doc.save(docx_path)
    print(f"Word document saved: {docx_path}")

    if os.path.exists(docx_path):
        print(f"File size: {os.path.getsize(docx_path)} bytes")
    else:
        print("File does not exist!")

    # PDF conversion is not available on macOS without Word application
    print("Note: PDF conversion is not available on macOS. Only Word document generated.")
    pdf_path = None

    return docx_path, pdf_path

def main():
    # Load the COA results Excel file
    # Note: On macOS, using local file instead of SharePoint path
    excel_file_path = os.path.join(os.getcwd(), "COA results.xlsx")
    
    try:
        coa_results = pd.read_excel(excel_file_path)
    except FileNotFoundError:
        print(f"Error: The file '{excel_file_path}' was not found.")
        print("Please ensure the file exists and the path is correct.")
        sys.exit(1)
    except Exception as e:
        print(f"An error occurred while trying to read the Excel file: {e}")
        sys.exit(1)

    # Set up the output directory
    output_directory = os.path.join(os.getcwd(), "COA Reports")
    os.makedirs(output_directory, exist_ok=True)

    # Set up the template path
    template_path = os.path.join(os.getcwd(), "COA TEMPLATE.docx")

    # Process each row in the Excel file
    for index, row in coa_results.iterrows():
        # Check if 'ToPrint' column exists and has value 'Y'
        if 'ToPrint' in row and row['ToPrint'] == 'Y':
            filled_doc = fill_template(row, template_path)
            
            # Create the new filename format
            release_date = format_date(row['Release Date'])
            brand = row['Brand']
            product = row['Product']
            flavor = row['Flavor']
            lot = row['Lot']
            
            # Remove any characters that are not allowed in filenames
            filename_base = f"{release_date}-{brand}-{product}-{flavor}-{lot}"
            filename_base = "".join(c for c in filename_base if c.isalnum() or c in ('-', '_')).rstrip()
            
            docx_path, pdf_path = save_as_docx_and_pdf(filled_doc, filename_base, output_directory)
            
            # Update the 'ToPrint' column to 'Success'
            coa_results.at[index, 'ToPrint'] = 'Success'
            
            print(f"Success: COA report generated for {filename_base}")

    # Save the updated Excel file
    try:
        coa_results.to_excel(excel_file_path, index=False)
        print(f"Excel file updated successfully: {excel_file_path}")
    except Exception as e:
        print(f"Error updating Excel file: {e}")

    print(f"COA reports have been successfully generated in {output_directory}")

if __name__ == "__main__":
    main()
